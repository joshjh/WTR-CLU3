import docx

class word_output:

    def __init__(self, docfile: object) -> object:
        """

        :rtype: docx document object
        """
        self.doc = docx.Document()
        self.docfile = docfile
        self.doc_header_int = 1

    def write_header(self, doc_header):
        self.doc.add_heading(doc_header, self.doc_header_int)

    def write_line(self, line):
        self.doc.add_paragraph(line)

    def doc_save(self):
        self.doc.save(self.docfile)


